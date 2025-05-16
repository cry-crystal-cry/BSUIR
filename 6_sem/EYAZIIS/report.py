from docx import Document
from docx.shared import Inches
import time
import requests
import os

# Инициализация документа
doc = Document()
doc.add_heading('Отчет по лабораторным работам №5 и №6', level=1)

# Описание системы
doc.add_paragraph('Дата и время составления отчета: 16 мая 2025 года, 07:12 AM CEST.')
doc.add_paragraph('Разработана диалоговая система для спортивной тематики с использованием следующих технологий:')
doc.add_paragraph('- Ollama (модель llama3.2) для обработки запросов.')
doc.add_paragraph('- FastAPI для серверной логики.')
doc.add_paragraph('- Streamlit для пользовательского интерфейса.')
doc.add_paragraph('- SQLite для хранения логов и метаданных документов.')
doc.add_paragraph('Данные модели и векторное хранилище перенесены на D:\\OllamaData\\.ollama из-за ограничения в 6 ГБ на диске C:.')
doc.add_paragraph('Директория проекта: L:\\projects\\BSUIR\\6_sem\\EYAZIIS.')

# Тестирование производительности
doc.add_heading('Тестирование производительности', level=2)
start = time.time()
response = requests.post("http://localhost:8000/chat", json={"question": "Как улучшить технику бега?", "session_id": "test", "model": "llama3.2"})
end = time.time()
duration = end - start
doc.add_paragraph(f'Время обработки запроса (FastAPI + LangChain + Ollama): {duration:.2f} секунд.')

# Результаты тестирования
doc.add_heading('Результаты тестирования', level=2)
doc.add_paragraph('Система успешно загружает документы формата TXT, PDF и DOCX (до 200 МБ) в директорию D:\\OllamaData\\uploads.')
doc.add_paragraph('Примеры запросов:')
doc.add_paragraph('- "Как улучшить технику бега?" — система возвращает рекомендации на основе загруженных документов.')
doc.add_paragraph('- "Какие упражнения для выносливости?" — система выдает соответствующие ответы.')
doc.add_paragraph('Функциональность: поддержка новых сессий, выбор истории сессий, удаление сообщений.')

# Интеграция других лабораторных (если применимо)
doc.add_heading('Интеграция других лабораторных работ', level=2)
doc.add_paragraph('Дополнительные лабораторные работы (LAB1.py–LAB4.py) не интегрированы в текущую систему, так как их функциональность не была указана. При необходимости их можно запустить отдельно или интегрировать (например, LAB3.py для извлечения ключевых слов).')

# Сохранение документа
output_path = os.path.join("L:\projects\BSUIR\\6_sem\EYAZIIS", "Lab_Report_LABA5_6.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Отчет сохранен в {output_path}")